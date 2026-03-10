# SPDX-License-Identifier: Apache-2.0
# Copyright (c) 2025–2026 Rohan R @rotsl

"""DOCX file loader for ContextFusion."""

from ..logging_utils import get_logger
from ..types import RawSegment
from .base_loader import BaseLoader

logger = get_logger("docx_loader")


class DocxLoader(BaseLoader):
    """Loader for DOCX files using python-docx."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def __init__(self):
        self._docx_available = None

    def _check_docx(self) -> bool:
        """Check if python-docx is available."""
        if self._docx_available is None:
            try:
                import docx

                self._docx_available = True
            except ImportError:
                self._docx_available = False
                logger.warning("python-docx not installed, DOCX loading disabled")
        return self._docx_available

    def supports(self, file_path: str) -> bool:
        """Check if this loader supports the given file.

        Args:
            file_path: Path to the file

        Returns:
            True if file is a DOCX
        """
        return self._get_extension(file_path) in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: str) -> list[RawSegment]:
        """Load segments from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of raw segments
        """
        if not self._check_docx():
            logger.error("python-docx required for DOCX loading")
            return []

        from docx import Document

        segments = []

        try:
            doc = Document(file_path)

            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    # Detect if it's a heading
                    style = para.style.name if para.style else "Normal"
                    is_heading = "Heading" in style

                    segment = RawSegment(
                        text=text,
                        metadata={
                            "paragraph_index": i,
                            "style": style,
                            "is_heading": is_heading,
                        },
                        source_path=file_path,
                    )
                    segments.append(segment)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))

                if table_text:
                    segment = RawSegment(
                        text="\n".join(table_text),
                        metadata={
                            "table_index": table_idx,
                            "is_table": True,
                        },
                        source_path=file_path,
                    )
                    segments.append(segment)

        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")

        return segments
